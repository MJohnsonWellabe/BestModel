"""Build the AM Best rating paper (.docx) for the ELT capital-appetite session.

Neutral, educational rewrite. Treats capital as one of two live levers (the coming RBC decline
is a real risk to the balance-sheet grade), teaches RBC and BCAR on their own terms, shows where
we sit versus peers, and lays out what realistically moves the rating with real examples.

Voice rules: plain language, no em-dashes, no decorative bolding, no "not X but Y" constructions.
Self-contained: builds its own figures and reads tool/data.json. Reuses the verified notching math.
Output: output/whitepaper/Wellabe_AMBest_Rating.docx
Run: python src/make_rating_paper.py
"""
from __future__ import annotations
import json
import re
import sys
from pathlib import Path
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
import notching  # noqa: E402

FIG = ROOT / "output" / "whitepaper" / "figures"
FIG.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "output" / "whitepaper" / "Wellabe_AMBest_Rating.docx"
# Prefer the full licensed frame when it has been built locally; otherwise use the published
# public frame, which carries the same public fields plus the BCAR scores from the Best reports.
_FRAME = ROOT / "tool" / "data.json"
if not _FRAME.exists():
    _FRAME = ROOT / "tool" / "public_data.json"
D = json.load(open(_FRAME))["carriers"]


def _bcar(c):
    return c.get("bcar_score")


def bcar_peers():
    """Carriers with a BCAR score, from their Best's Credit Reports."""
    return [c for c in D if _bcar(c) is not None]


def _corr(xs, ys):
    n = len(xs)
    if n < 3:
        return float("nan")
    mx, my = sum(xs) / n, sum(ys) / n
    num = sum((x - mx) * (y - my) for x, y in zip(xs, ys))
    dx = (sum((x - mx) ** 2 for x in xs)) ** 0.5
    dy = (sum((y - my) ** 2 for y in ys)) ** 0.5
    return num / (dx * dy) if dx and dy else float("nan")


def _median(xs):
    xs = sorted(x for x in xs if x is not None)
    if not xs:
        return None
    n = len(xs)
    return xs[n // 2] if n % 2 else (xs[n // 2 - 1] + xs[n // 2]) / 2


# Facts computed from the reports so the prose cannot drift from the data.
_BOTH = [c for c in bcar_peers() if c.get("rbc_cal_pct") is not None]
_N_BOTH = len(_BOTH)
_R_BCAR_RBC = _corr([c["rbc_cal_pct"] for c in _BOTH], [_bcar(c) for c in _BOTH])
_TIER_MED = {t: _median([_bcar(c) for c in bcar_peers() if c.get("bs_assessment") == t])
             for t in ("Strongest", "Very Strong", "Strong", "Adequate")}
_RANKED = sorted(bcar_peers(), key=lambda c: -_bcar(c))
_W_RANK = next((i + 1 for i, c in enumerate(_RANKED) if c.get("is_wellabe")), None)

ACCENT = "#2E5A88"; INK = "#0B1C2C"; WELL = "#C0392B"; MUTE = "#5C6B78"
TIERCOL = {"Strongest": "#0F5C8C", "Very Strong": "#27A35A", "Strong": "#9AA0A6", "Adequate": "#C97B2B"}
plt.rcParams.update({"font.size": 11, "font.family": "DejaVu Sans", "axes.edgecolor": "#C9BFA8",
                     "axes.titlesize": 13, "axes.titleweight": "bold", "figure.dpi": 150})

FLOORS = {"Strongest": 530, "Very Strong": 375, "Strong": 275}


def Wc():
    return next(c for c in D if c.get("is_wellabe"))


# ============================================================ figures
def fig_bcar_history():
    yrs = ["2022", "2023", "2024", "2025"]; bcar = [73.4, 73.0, 71.2, 67.3]
    fig, ax = plt.subplots(figsize=(7.0, 3.5), constrained_layout=True)
    ax.bar(yrs, bcar, color=ACCENT, width=0.6, zorder=3)
    for x, v in zip(yrs, bcar):
        ax.text(x, v + 1.6, f"{v:.0f}%", ha="center", fontweight="bold", color=INK)
    ax.annotate("plan takes it\nlower from here", xy=(3.35, 40), xytext=(3.35, 58), color="#8A8A8A",
                fontsize=8.5, ha="center", va="top", arrowprops=dict(arrowstyle="-|>", color="#8A8A8A", lw=1.4))
    ax.axhspan(0, 25, color=WELL, alpha=0.08, zorder=0)
    ax.axhline(25, color=WELL, lw=1.4, ls="--", zorder=2)
    ax.text(3.4, 28, "25% is the line for the top tier", color=WELL, ha="right", fontsize=10)
    ax.set_ylim(0, 88); ax.set_ylabel("BCAR cushion at the 1-in-250 stress (%)")
    ax.set_title("BCAR over the four reviews we have: a wide cushion, drifting down")
    ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "rp_bcar.png"); plt.close(fig)


def fig_rbc_path():
    yrs = [2025, 2026, 2027, 2028, 2029, 2030]
    rbc = [648, 585, 510, 445, 400, 420]   # illustrative path to the ~400% plan trough around 2029
    fig, ax = plt.subplots(figsize=(7.0, 3.8), constrained_layout=True)
    ax.axhspan(FLOORS["Strongest"], 720, color=TIERCOL["Strongest"], alpha=0.10)
    ax.axhspan(FLOORS["Very Strong"], FLOORS["Strongest"], color=TIERCOL["Very Strong"], alpha=0.10)
    ax.axhspan(FLOORS["Strong"], FLOORS["Very Strong"], color=TIERCOL["Strong"], alpha=0.12)
    for name, y in FLOORS.items():
        ax.axhline(y, color="#9AA0A6", ls=":", lw=1)
        ax.text(2030.2, y, f"  {name} floor ~{y}%", va="center", fontsize=8.5, color=MUTE)
    ax.plot(yrs, rbc, color=ACCENT, lw=2.2, marker="o", ms=5, zorder=4)
    ax.scatter([2025], [648], color=WELL, s=90, zorder=5)
    ax.annotate("today ~648%, Strongest", (2025, 648), xytext=(4, 8), textcoords="offset points",
                color=WELL, fontweight="bold", fontsize=9)
    ax.annotate("plan trough ~400%,\nin the Very Strong band", (2029, 400), xytext=(-2, -36),
                textcoords="offset points", color=INK, fontsize=9, ha="center")
    ax.set_ylim(250, 720); ax.set_xlim(2024.8, 2031.5)
    ax.set_ylabel("NAIC RBC ratio, CAL basis (%)")
    ax.set_title("Where the plan takes our RBC ratio (illustrative)")
    ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "rp_rbcpath.png"); plt.close(fig)


def _box_by_tier(ax, groups, key, gkey, log=False):
    for c in D:
        c["_g"] = c.get(gkey)
    data = [[c[key] for c in D if c.get("_g") == g and c.get(key) is not None] for g in groups]
    pos = list(range(len(groups)))
    bp = ax.boxplot(data, positions=pos, orientation="horizontal", widths=0.55, patch_artist=True,
                    showfliers=False, medianprops=dict(color=INK, lw=1.6),
                    whiskerprops=dict(color="#8A8A8A"), capprops=dict(color="#8A8A8A"))
    for patch, g in zip(bp["boxes"], groups):
        patch.set_facecolor(TIERCOL.get(g, "#9AA0A6")); patch.set_alpha(0.35); patch.set_edgecolor("#8A8A8A")
    for i, ys in enumerate(data):
        ax.scatter(ys, [i] * len(ys), color=TIERCOL.get(groups[i], "#9AA0A6"), s=18, alpha=0.55,
                   zorder=3, edgecolor="white", lw=.4)
    ax.set_yticks(pos); ax.set_yticklabels(groups)
    if log:
        ax.set_xscale("log")
    return data


def fig_cap_tiers():
    order = ["Strongest", "Very Strong", "Strong", "Adequate"]
    fig, ax = plt.subplots(figsize=(7.0, 3.6), constrained_layout=True)
    _box_by_tier(ax, order, "rbc_cal_pct", "bs_assessment", log=True)
    w = Wc()
    ax.scatter([w["rbc_cal_pct"]], [0], marker="D", s=160, color=WELL, edgecolor=INK, lw=1.4, zorder=6)
    ax.annotate("Wellabe, ~648%", (w["rbc_cal_pct"], 0), xytext=(6, 15), textcoords="offset points",
                color=WELL, fontweight="bold", fontsize=9.5)
    ax.set_xticks([200, 300, 500, 1000, 2000]); ax.set_xticklabels(["200", "300", "500", "1,000", "2,000"])
    ax.set_xlabel("NAIC RBC ratio, CAL basis (%)")
    ax.set_title("Capital by balance-sheet tier across peers")
    ax.invert_yaxis(); ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "rp_captiers.png"); plt.close(fig)


def fig_bcar_tiers():
    """BCAR by balance-sheet tier: the measure Best actually grades on, across the peer set."""
    order = ["Strongest", "Very Strong", "Strong", "Adequate"]
    peers = bcar_peers()
    fig, ax = plt.subplots(figsize=(7.0, 3.6), constrained_layout=True)
    data = [[_bcar(c) for c in peers if c.get("bs_assessment") == g] for g in order]
    pos = list(range(len(order)))
    bp = ax.boxplot(data, positions=pos, orientation="horizontal", widths=0.55, patch_artist=True,
                    showfliers=False, medianprops=dict(color=INK, lw=1.6),
                    whiskerprops=dict(color="#8A8A8A"), capprops=dict(color="#8A8A8A"))
    for patch, g in zip(bp["boxes"], order):
        patch.set_facecolor(TIERCOL.get(g, "#9AA0A6")); patch.set_alpha(0.35); patch.set_edgecolor("#8A8A8A")
    for i, ys in enumerate(data):
        ax.scatter(ys, [i] * len(ys), color=TIERCOL.get(order[i], "#9AA0A6"), s=18, alpha=0.55,
                   zorder=3, edgecolor="white", lw=.4)
    ax.set_yticks(pos); ax.set_yticklabels(order)
    w = Wc()
    if _bcar(w) is not None:
        ax.scatter([_bcar(w)], [0], marker="D", s=160, color=WELL, edgecolor=INK, lw=1.4, zorder=6)
        ax.annotate(f"Wellabe, {_bcar(w):.1f}%", (_bcar(w), 0), xytext=(6, 15), textcoords="offset points",
                    color=WELL, fontweight="bold", fontsize=9.5)
    ax.axvline(25, color=WELL, lw=1.4, ls="--", zorder=2)
    ax.text(25, 3.55, " 25% top-tier line", color=WELL, fontsize=9, va="top")
    ax.set_xlabel("BCAR cushion at the 1-in-250 stress (%)")
    ax.set_title("BCAR by balance-sheet tier: stronger tiers sit higher, but they overlap")
    ax.invert_yaxis(); ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "rp_bcartiers.png"); plt.close(fig)


def fig_bcar_vs_rbc():
    """The two capital yardsticks against each other: they do not move together."""
    peers = [c for c in bcar_peers() if c.get("rbc_cal_pct") is not None]
    fig, ax = plt.subplots(figsize=(7.0, 4.0), constrained_layout=True)
    for c in peers:
        grp = (c.get("rating_basis") == "group-member")
        ax.scatter([c["rbc_cal_pct"]], [_bcar(c)], s=42, alpha=.75, zorder=3,
                   color="#C97B2B" if grp else ACCENT, edgecolor="white", lw=.5)
    w = Wc()
    ax.scatter([w["rbc_cal_pct"]], [_bcar(w)], marker="D", s=170, color=WELL, edgecolor=INK, lw=1.4, zorder=6)
    ax.annotate(f"Wellabe\n{w['rbc_cal_pct']:.0f}% RBC, {_bcar(w):.1f}% BCAR", (w["rbc_cal_pct"], _bcar(w)),
                xytext=(-10, -42), textcoords="offset points", color=WELL, fontweight="bold",
                fontsize=9, ha="center")
    from matplotlib.lines import Line2D
    ax.legend(handles=[Line2D([], [], marker="o", ls="", color=ACCENT, label="rated on its own"),
                       Line2D([], [], marker="o", ls="", color="#C97B2B", label="rated with group support")],
              frameon=False, fontsize=9, loc="upper left")
    ax.set_xscale("log")
    ax.set_xlim(120, 2600); ax.set_ylim(-70, 92)
    ax.axhspan(-70, 25, color=WELL, alpha=0.06, zorder=0)
    ax.axhline(25, color=WELL, lw=1.3, ls="--", zorder=2)
    ax.text(2500, 27, "25% top-tier line on BCAR ", color=WELL, fontsize=9, ha="right")
    ax.set_xticks([200, 300, 500, 1000, 2000]); ax.set_xticklabels(["200", "300", "500", "1,000", "2,000"])
    ax.set_xlabel("NAIC RBC ratio, CAL basis (%), log scale")
    ax.set_ylabel("BCAR cushion at the 1-in-250 stress (%)")
    ax.set_title("The two capital measures do not move together")
    ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "rp_bcarrbc.png"); plt.close(fig)


def fig_earn_tiers():
    order = ["Strong", "Adequate", "Marginal"]
    fig, ax = plt.subplots(figsize=(7.0, 3.6), constrained_layout=True)
    _box_by_tier(ax, order, "roe_5yr_mean", "op_assessment")
    w = Wc()
    ax.scatter([w["roe_5yr_mean"]], [1], marker="D", s=160, color=WELL, edgecolor=INK, lw=1.4, zorder=6)
    ax.annotate("Wellabe: losing money, held at Adequate", (w["roe_5yr_mean"], 1),
                xytext=(8, -30), textcoords="offset points", color=WELL, fontweight="bold", fontsize=9.5)
    for nm, lab in [("Pekin Life", "Pekin (A-)"), ("Government Personnel Mutual", "GPM (B++)")]:
        c = next((x for x in D if x["rating_unit_name"] == nm), None)
        if c and c.get("roe_5yr_mean") is not None:
            ax.annotate(lab, (c["roe_5yr_mean"], 2), xytext=(0, 10), textcoords="offset points",
                        ha="center", fontsize=9, color="#7A2418")
    ax.axvline(0, color="#C9BFA8", lw=1)
    ax.set_xlabel("Five-year average return on equity (%)")
    ax.set_title("Operating performance by tier across peers")
    ax.invert_yaxis(); ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "rp_earntiers.png"); plt.close(fig)


def fig_ladder():
    fsr = ["A", "A-", "B++"]; xi = {f: i for i, f in enumerate(fsr)}
    rows = [
        ("Today", "A", ACCENT, "Strongest capital, Adequate operating"),
        ("Capital slips one tier (RBC near the trough)", "A-", "#27A35A", "one letter, like Globe Life or American Southern"),
        ("Operating slips to Marginal", "A-", "#27A35A", "one letter, the Strongest grade holds it"),
        ("Both slip together", "B++", WELL, "the real tail, like GPM today"),
    ]
    fig, ax = plt.subplots(figsize=(7.0, 3.5), constrained_layout=True)
    for i, (lab, f, col, note) in enumerate(rows):
        ax.barh(i, xi[f] + 0.5, color=col, alpha=.85, zorder=3, height=0.6)
        ax.text(xi[f] + 0.58, i, f, va="center", fontweight="bold", color=INK, fontsize=12)
        ax.text(xi[f] + 0.95, i, note, va="center", color=MUTE, fontsize=8.5)
        ax.text(-0.08, i, lab, va="center", ha="right", fontsize=9.5)
    ax.set_xlim(0, 5.4); ax.set_yticks([]); ax.invert_yaxis()
    ax.set_xticks(range(len(fsr))); ax.set_xticklabels(fsr)
    ax.set_title("How far we could fall, and what each step takes")
    ax.set_xlabel("Financial strength rating")
    for s in ["top", "right", "left"]:
        ax.spines[s].set_visible(False)
    fig.savefig(FIG / "rp_ladder.png"); plt.close(fig)


def fig_msstress():
    fig = plt.figure(figsize=(8.4, 2.7))
    ax = fig.add_axes([0, 0, 1, 1]); ax.axis("off"); ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.text(0.5, 0.91, "How Med Supp concentration would actually move us: through earnings",
            ha="center", fontsize=12.5, fontweight="bold", color=INK)
    chain = [(0.13, "#FBE9E7", WELL, "Concentrated\nMed Supp book"),
             (0.38, "#FBE9E7", WELL, "Shock: rate,\nreg, or competitor"),
             (0.63, "#EEF2F6", ACCENT, "Loss-ratio\nvolatility"),
             (0.88, "#EEF2F6", ACCENT, "Operating slips,\nrating to A-")]
    bw, bh, by = 0.205, 0.23, 0.50
    for i, (x, fc, ec, txt) in enumerate(chain):
        ax.add_patch(plt.Rectangle((x - bw / 2, by), bw, bh, fc=fc, ec=ec, lw=1.3))
        ax.text(x, by + bh / 2, txt, ha="center", va="center", fontsize=8.7, color=ec, fontweight="bold")
        if i < len(chain) - 1:
            ax.annotate("", xy=(chain[i + 1][0] - bw / 2 - 0.004, by + bh / 2),
                        xytext=(x + bw / 2 + 0.004, by + bh / 2),
                        arrowprops=dict(arrowstyle="-|>", color="#8A8A8A", lw=1.7))
    ax.text(0.5, 0.34, "Franchise grade holds at Neutral. Capital is touched only indirectly, through weaker earnings.",
            ha="center", fontsize=8.6, color=MUTE, style="italic")
    ax.text(0.5, 0.18, "Same shape as Pekin: a concentrated book, a shock, then a Negative outlook from earnings, not a franchise cut.",
            ha="center", fontsize=8.6, color=MUTE, style="italic")
    fig.savefig(FIG / "rp_msstress.png", dpi=150); plt.close(fig)


# ============================================================ docx helpers
def _cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def part(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(2)


def H(doc, text, before=12):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(13.5)
    r.font.color.rgb = RGBColor.from_string(ACCENT.lstrip("#"))
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(2)


def body(doc, text, size=10.5):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6); return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(text); r.font.size = Pt(10.5); return p


_FIGN = [0]


def img(doc, name, caption):
    """Insert a figure. Any leading "Figure N." in the caption is replaced with the running
    number so captions stay in order when figures are added or moved."""
    _FIGN[0] += 1
    caption = re.sub(r"^\s*Figure\s+\d+\.\s*", "", caption)
    caption = f"Figure {_FIGN[0]}. {caption}"
    doc.add_picture(str(FIG / name), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTE.lstrip("#"))
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(10)


def table(doc, headers, rows, highlight=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("FFFFFF"); _cell_bg(c, ACCENT.lstrip("#"))
    for row in rows:
        cells = t.add_row().cells
        hot = highlight and row[0] == highlight
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
            if hot:
                r.bold = True; r.font.color.rgb = RGBColor.from_string(WELL.lstrip("#"))
                _cell_bg(cells[j], "FBE9E7")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================================================ peer table
CURATED = [
    "Guardian Life", "Physicians Mutual", "Mutual of Omaha", "Aflac",
    "Guarantee Trust Life", "National Guardian Life", "CNO — Bankers Life", "American National",
    "Globe Life", "American-Amicable / Trinity", "Wellabe Group",
    "National Western Life", "Assurity Life", "Pekin Life", "Funeral Directors Life",
    "Homesteaders Life", "Atlantic American — American Southern",
    "ManhattanLife — Assurance", "Government Personnel Mutual", "Investors Heritage",
    "ManhattanLife — Western United", "Continental General", "Sentinel Security Life",
]
FSR_ORDER = ["A++", "A+", "A", "A-", "B++", "B+", "B", "B-"]


def peer_rows():
    by = {c["rating_unit_name"]: c for c in D}
    out = []
    for nm in CURATED:
        c = by.get(nm)
        if not c:
            continue
        rbc = c.get("rbc_cal_pct")
        bc = _bcar(c)
        out.append([nm.replace("Wellabe Group", "Wellabe").replace(" — ", ", "),
                    c.get("fsr") or "n/a", c.get("bs_assessment") or "n/a",
                    c.get("op_assessment") or "n/a", c.get("bp_assessment") or "n/a",
                    c.get("erm_assessment") or "n/a",
                    f"{bc:.1f}%" if bc is not None else "n/a",
                    f"{round(rbc):,}%" if rbc else "n/a"])
    out.sort(key=lambda r: (FSR_ORDER.index(r[1]) if r[1] in FSR_ORDER else 99,
                            -float(r[7].replace(",", "").rstrip("%")) if r[7] != "n/a" else 0))
    return out


# ============================================================ build
def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

    t = doc.add_paragraph(); r = t.add_run("Our AM Best Rating: How It Works, Where We Stand, and What Could Move It")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    s = doc.add_paragraph(); rs = s.add_run("Office of Strategy. ELT, internal and confidential.")
    rs.italic = True; rs.font.size = Pt(10); rs.font.color.rgb = RGBColor.from_string(MUTE.lstrip("#"))

    H(doc, "What this paper is for", before=10)
    body(doc, "A key component of our company-wide risk appetite is to maintain an AM Best rating of A- or better. "
              "This paper explains how the rating is built, where we stand, and what would move us, so the "
              "conversation runs off shared mechanics rather than impressions.")
    body(doc, "One fact drives the rest of it. Our plan draws capital down over the next few years. The RBC ratio "
              "falls from around 648% at 12/31/25 toward a trough near 400% around 2029 before it rebuilds, and that "
              "trough already assumes the planned surplus note and a C-3 Phase I reserve change. The decline is "
              "deliberate. The question this paper answers is whether it threatens the rating.")

    # ---------- PART I
    part(doc, "Part I.  How the rating works")

    H(doc, "1.  What the rating is, and where it matters")
    body(doc, "The financial strength rating is Best's opinion of our ability to pay claims. It is a business asset, "
              "and it earns its keep in specific places. It matters in preneed, where funeral-home programs place "
              "obligations that run for decades and prefer a stable, highly rated counterparty. It matters in "
              "reinsurance and in any capital raise, where the letter sets our terms. It matters least in "
              "agent-sold Medicare Supplement, where the product is short-duration and backed by state guaranty "
              "funds, so the buyer is insulated from our balance sheet.")
    body(doc, "The line that matters is A-. At A- and above we stay on the shelf almost everywhere. Below A-, "
              "distribution starts closing: shelf access, preneed programs and reinsurance terms are the first to go. "
              "That is why the appetite is set at A- rather than at our current A.")

    H(doc, "2.  What we looked at")
    body(doc, f"To read our own rating against the market we pulled the Best credit report for {len(bcar_peers())} "
              "carriers: our direct competitors, the senior-market and preneed specialists we are compared to, and a "
              "set of large diversified carriers for range. They span roughly 100 million dollars to 320 billion "
              "dollars in assets and every rating from A++ to B. Each report carries the four building-block grades "
              "and, importantly, the BCAR score, which is published nowhere else.")
    body(doc, "That gives us the same inputs Best used, for us and for the peer set, so the comparisons in this paper "
              "are like for like rather than inferred from financial statements.")

    H(doc, "3.  The two capital yardsticks: RBC and BCAR")
    body(doc, "Two capital measures get used for two different jobs, and conflating them causes most of the confusion "
              "in these discussions.")
    body(doc, "NAIC risk-based capital, which we report on the CAL basis and which runs about 648%, is a regulatory "
              "floor. It compares our total adjusted capital to a control level the regulator calculates, and it "
              "answers one question: is the regulator comfortable. Once a carrier is well clear of the floor, the "
              "ratio says very little about relative strength, and Best does not use it to set the rating.")
    table(doc, ["Regulatory level (CAL basis)", "Ratio", "What happens"],
          [["Company Action Level", "below 100%", "File a corrective plan"],
           ["Regulatory Action Level", "below 75%", "Regulator prescribes action"],
           ["Authorized Control Level", "below 50%", "Regulator may take control"],
           ["Mandatory Control Level", "below 35%", "Regulator must take control"]])
    body(doc, "BCAR is Best's own capital model, and it is the one that feeds the rating. Best runs our balance sheet "
              "through a stress up to roughly a 1-in-250-year loss and measures how much capital is still standing. "
              "A score above 25% at that stress is the level associated with the top capital assessment. BCAR picks "
              "up asset risk, reserve adequacy and catastrophe exposure that the RBC formula treats more coarsely.")
    body(doc, f"The reports show how far apart the two run. Across the {_N_BOTH} carriers with both measures the "
              f"correlation is {_R_BCAR_RBC:+.2f}, which is close to no relationship. Aetna's American Continental "
              "carries an RBC ratio above 2,000% on a BCAR of 29%. ManhattanLife Assurance runs 859% RBC on a BCAR "
              "of 17.7%. A high RBC ratio does not buy a strong capital assessment, which is why Best built its own "
              "model.")
    img(doc, "rp_bcarrbc.png", "Best's capital model against the regulator's ratio for "
        f"{_N_BOTH} carriers. The two do not move together. Orange carriers are rated with group support and "
        "cluster low on BCAR: the parent's capital, not their own, holds the rating up.")
    body(doc, "So we track two numbers for two purposes. RBC is the regulatory floor and the number our plan moves "
              "most. BCAR is the number that feeds our grade.")

    H(doc, "4.  The balance-sheet grade starts with BCAR and does not end there")
    body(doc, "The balance-sheet grade sets the starting point for the whole rating:")
    table(doc, ["Balance-sheet grade", "Starting point (ICR)", "Letter it implies"],
          [["Strongest", "a+ / a", "A"], ["Very Strong", "a / a-", "A / A-"],
           ["Strong", "a- / bbb+", "A- / B++"], ["Adequate", "bbb+ / bbb / bbb-", "B++ / B+"]])
    body(doc, "BCAR is the largest input to that grade, and across the peer set the median score steps down cleanly "
              f"with each tier: about {_TIER_MED['Strongest']:.0f}% for Strongest, {_TIER_MED['Very Strong']:.0f}% "
              f"for Very Strong, {_TIER_MED['Strong']:.0f}% for Strong and {_TIER_MED['Adequate']:.0f}% for "
              "Adequate. The tiers overlap at the edges, so the score indicates the grade without determining it.")
    img(doc, "rp_bcartiers.png", "BCAR by balance-sheet tier across the carriers we hold reports for. The tiers "
        "step down in the expected order and overlap at the edges.")
    body(doc, "Clearing 25% does not earn the top grade. The threshold applies to risk-adjusted capitalization, "
              "which is one component of the balance-sheet assessment. In 18 of the 44 reports Best states that the "
              "carrier's risk-adjusted capitalization sits at the strongest level. Thirteen of those 18 are graded "
              "below Strongest on the overall balance sheet.")
    body(doc, "The reasons are visible in the same paragraphs, and they are mostly capital that will not stay put. "
              "Reserve National's parent anticipates upstreaming dividends. National Western's balance sheet is "
              "expected to be affected by planned capital withdrawals tied to its acquisition. Others carry "
              "concentrated reserves or heavy allocations to lower-quality bonds. Best also weighs the absolute size "
              "of the capital base, liquidity, reserve adequacy and financial flexibility.")
    body(doc, "Two of those factors deserve attention because they are the levers we are pulling. On reinsurance, the "
              "reports do not support the intuition that heavy use costs a carrier the grade. Best describes it in "
              "neutral or favourable terms across the peer set: appropriate use to relieve new-sales strain at GPM, "
              "risk lessened through reinsurance at National Western, usage still below the industry average at "
              "Liberty Bankers. Where it draws a comment, as with American-Amicable's material use relative to peers, "
              "Best immediately notes the cessions go to high-rated counterparties. Reinsurance is read as risk "
              "management, and the quality of the counterparty is what it looks at.")
    body(doc, "Debt is treated differently, and Best is explicit about it. National Guardian Life carries surplus "
              "notes and Best writes that they slightly decrease the quality of capital overall, while noting the "
              "group is paying them down. That carrier is graded Very Strong rather than Strongest. The penalty is "
              "calibrated rather than automatic: Guardian Life holds the Strongest grade with surplus notes "
              "outstanding, because Best judges its financial leverage favourable and its interest coverage within "
              "tolerance. Where leverage is heavier and sits alongside other issues, as at ManhattanLife, it "
              "contributes to a materially lower grade. Section 8 returns to what this means for our own surplus "
              "note.")

    H(doc, "5.  The other three blocks, and which of them can move us")
    body(doc, "The balance sheet sets the starting point. The other three blocks add or remove notches from it:")
    table(doc, ["Building block", "How far it can move a rating", "No-change grade"],
          [["Operating performance (earnings)", "up 2, down 3", "Adequate"],
           ["Business profile (franchise)", "up 2, down 2", "Neutral"],
           ["Risk management", "up 1, down 4", "Appropriate"]])
    body(doc, "The downside on each block is larger than the upside, so ground is easier to lose than to gain. Of "
              "the three, only operating performance is realistically in play for us.")
    body(doc, "Business profile is Best's read of our franchise: scale, market position, product and geographic "
              "diversification, and distribution. We are graded Neutral. Moving up requires materially greater scale, "
              "which is a multi-year strategic change rather than a plan variance. It stays fixed unless we change "
              "the shape of the company ourselves, and section 7 covers the one decision that would.")
    body(doc, "Risk management is graded Appropriate, and it is effectively fixed. The upside is essentially "
              "unavailable, since Very Strong is reserved for programs well beyond our size and complexity. The "
              "grade rarely moves a rating on its own. The exposure here is not the grade, it is allowing a "
              "preventable event, a failed reinsurance recovery or a governance lapse, to reframe how Best reads "
              "everything else.")
    body(doc, "That leaves the balance sheet and operating performance as the two live blocks, and Part III returns "
              "to them.")

    # ---------- PART II
    part(doc, "Part II.  Where we stand")

    H(doc, "6.  Our grade today")
    body(doc, "Best grades us Strongest on the balance sheet, Adequate on operating performance, Neutral on business "
              "profile and Appropriate on risk management. Strongest opens us at a, the other three blocks are all at "
              "no change, and we land at A with a Stable outlook.")
    body(doc, f"Our BCAR is 67.3%, which ranks {_W_RANK} of the {len(bcar_peers())} carriers we hold reports for. "
              "None of the qualifiers from section 4 currently applies to us: there is no parent upstreaming our "
              "capital, no acquisition-driven withdrawal, and no reliance on reinsurance to carry new-business "
              "strain. That is why the capital assessment converts cleanly into the top grade rather than being "
              "marked down the way it was for thirteen of our peers.")
    table(doc, ["Year", "Net income", "Capital and surplus", "A&H combined ratio", "BCAR"],
          [["2021", "+$22M", "$630M", "95%", "n/a"], ["2022", "-$1M", "$615M", "99%", "73.4%"],
           ["2023", "-$21M", "$602M", "103%", "73.0%"], ["2024", "-$52M", "$560M", "109%", "71.2%"],
           ["2025", "-$71M", "$531M", "116%", "67.3%"]])
    img(doc, "rp_bcar.png", "BCAR across the four annual reviews we have. A wide cushion above the 25% level, "
        "trending down as losses draw on surplus.")
    body(doc, "Operating performance is the soft block. We have run losses for four straight years and they have "
              "grown each year, driven by the accident-and-health combined ratio moving from 95% to 116%. Our "
              "five-year return on equity sits near the bottom of the peer set. Best holds us at Adequate rather "
              "than Marginal because it reads the losses as the cost of funding growth in a maturing Medicare "
              "Supplement block, and because it accepts the plan that turns them. That grade rests on a forecast, "
              "which is the exposure.")
    body(doc, "The turn has started to show. Through the first half of 2026 both the loss ratio and the combined "
              "ratio are down about four points against 2025. That is the first hard evidence supporting the "
              "forecast Best has extended to us, and it is the number to keep producing.")

    H(doc, "7.  How we compare")
    body(doc, "Reading the same grades across the peer set shows where letters actually come from. The table is "
              "sorted by rating, then by capital.")
    table(doc, ["Carrier", "Rating", "RBC (CAL)", "BCAR", "Balance sheet", "Operating", "Business profile", "ERM"],
          peer_rows(), highlight="Wellabe")
    body(doc, "Capital does not sort the ratings. Globe Life holds our same A on a Strong balance sheet, a 316% RBC "
              "ratio and a BCAR of 6.6%, carried by scale and business profile. Guarantee Trust Life holds an A above "
              "800% RBC. What sorts the ratings is operating performance and business profile.")
    body(doc, "The low-capital examples come with a condition attached. Combined Insurance is rated A+ on a BCAR of "
              "2.7%, the three CNO companies are rated A on 3.8%, Forethought is rated A on 4.9%, and Humana is "
              "rated A on a BCAR below zero. Every one of them is rated with group support, so the parent's balance "
              "sheet is holding the letter up. The standalone exceptions are carriers of a size we will not reach, "
              "New York Life at 259 billion dollars in assets and Lincoln at 320 billion. We are rated standalone at "
              "2.3 billion, so our own capital has to do the work.")
    body(doc, "The peer set also answers a question worth settling, which is whether exiting a product line such as "
              "preneed would cost us the franchise grade through lost diversification. On the evidence, "
              "diversification is not what Best grades. Across the peer set the correlation between segment "
              "concentration and the business-profile grade is about zero, while the correlation with scale is "
              "strong. Mutual of Omaha is effectively monoline and graded Favorable. Assurity is as diversified as we "
              "are and graded Limited, because, in Best's words, it lacks significant market share. Funeral Directors "
              "writes preneed only and holds Neutral, because its business model cannot be easily replicated. Market "
              "position is what separates the tiers.")
    body(doc, "Our own report, though, names preneed specifically. Best writes that our Medicare Supplement "
              "concentration is somewhat offset by the diversification added through preneed, hospital indemnity, "
              "short-term care, and life and annuity. Exiting preneed would remove a stated mitigant and leave the "
              "competitive characterisation of Medicare Supplement standing on its own, at a point where we are "
              "already near the size boundary between the Neutral and Limited groups. The reassuring precedent is "
              "Physicians Mutual, which is Medicare Supplement concentrated at a similar size and holds Neutral on "
              "the strength of established market position and diversified distribution. A preneed exit is therefore "
              "defensible on the franchise grade, but it is not free, and it would need to be presented to Best "
              "alongside the market position and distribution that replace the offset.")
    body(doc, "We are an unusual profile in this set: top-decile capital paired with bottom-decile earnings. The "
              "typical A-rated peer is the reverse, a Very Strong balance sheet with Strong operating performance. "
              "We hold our A on capital and a plan Best believes, where others hold it on profits. That is the "
              "position the drawdown now tests.")

    # ---------- PART III
    part(doc, "Part III.  What realistically moves us")

    H(doc, "8.  The most likely outcome")
    body(doc, "The base case is that we stay at A. Capital falls as planned and stays well above the level "
              "associated with the top assessment, the operating turn continues to show in the combined ratio, and "
              "Best keeps both grades where they are. The first half of 2026 is consistent with that path.")
    body(doc, "The most likely adverse outcome is a single-notch move to A-, and it can arrive through either live "
              "block. If the balance-sheet grade slips from Strongest to Very Strong, our starting point moves from "
              "a to a- and the letter follows to A-. If operating performance slips from Adequate to Marginal, we "
              "lose a notch from the same starting point and land in the same place. Either one costs one letter. "
              "At A- we remain above the appetite line and on the shelf in every channel that matters.")
    body(doc, "A slip in the balance-sheet grade is the more likely of the two, and not because of the RBC trough. "
              "No RBC level sorts the grade. The risk is that a multi-year, deliberate capital drawdown is "
              "recognisably the same category of fact that cost thirteen peers the top tier: capital leaving the "
              "balance sheet.")
    body(doc, "The surplus note belongs in the same assessment, and the peer evidence tells us how it will be read. "
              "It raises the RBC ratio and BCAR while adding capital that Best treats as lower quality, the effect "
              "it described at National Guardian Life as slightly decreasing the quality of capital overall. Our "
              "current report says the group has no debt and minimal operating leverage, so the note replaces a "
              "clean statement with a qualified one, in the same window as the drawdown. It should not cost us the "
              "grade by itself. Guardian Life holds the Strongest assessment with surplus notes outstanding because "
              "Best judges its interest coverage to be within tolerance and its leverage favourable. That is the "
              "standard to hold ourselves to: coverage inside Best's tolerances and a visible paydown path, "
              "presented before Best has to ask.")

    H(doc, "9.  The catastrophic outcome")
    body(doc, "The downside case is both blocks moving in the same window. A balance-sheet slip to Very Strong plus "
              "operating performance to Marginal costs two letters and takes us to B++, below the appetite line.")
    img(doc, "rp_ladder.png", "Each lever on its own costs one letter. B++ requires both at the same time.")
    body(doc, "The two are stressed in the same window, which is what makes this more than arithmetic. The plan "
              "draws capital to its trough around 2029, which is also when the operating turn is most under test. A "
              "scenario where the Medicare Supplement turn stalls while capital is at its lowest is the one that "
              "produces both moves together.")
    body(doc, "At B++ the damage is concentrated where the rating actually sells. Preneed programs and funeral-home "
              "relationships come under immediate pressure, reinsurance terms tighten, and any capital raise gets "
              "materially more expensive. Guaranty-backed agent Medicare Supplement is the most insulated.")
    body(doc, "What happens next matters more than the drop itself, because B++ is a trough rather than a resting "
              "place. The operating slip that takes us there is a move from Adequate to Marginal, and Marginal means "
              "Best has stopped believing the losses turn. If profitability arrives shortly after, which is what the "
              "plan and the current combined-ratio trend both point to, that block returns to Adequate. Adequate is "
              "the no-change state, so recovering it does not require us to earn a positive notch. It only requires "
              "us to stop losing one. On a Very Strong balance sheet that puts us back at a- and A-, inside the "
              "appetite, and if capital rebuilds past the trough as planned the starting point returns to a and the "
              "letter to A.")
    body(doc, "The realistic shape of the downside is therefore a dip to B++ followed by a recovery to A-, rather "
              "than a permanent reset. Two things decide whether it plays out that way. The first is how long we sit "
              "at B++, because Best upgrades on demonstrated results rather than forecasts, so the recovery lags the "
              "earnings by at least a review cycle. The second is what happens to distribution while we are there. "
              "Carriers that fall below A- and lose shelf access can find the earnings recovery undercut by the "
              "business they lost, and that is the mechanism that keeps some of them down. Our exposure to it runs "
              "through preneed and reinsurance rather than through guaranty-backed Medicare Supplement, so holding "
              "those relationships through the trough is what converts a dip into a full recovery.")
    body(doc, "We assess the two-block scenario as a low-probability path. It requires both blocks to fail together "
              "rather than either one alone, we would have warning before it arrived, and the recovery route back to "
              "A- is open provided distribution holds.")

    H(doc, "10.  What we monitor, and how we tell the story")
    body(doc, "Four things carry the signal, and we should be deliberate about each.")
    bullet(doc, "The BCAR trajectory, because it feeds the balance-sheet grade. The 2026 report is the informative "
                "one: it will carry the surplus note and a further year of losses, and it tells us the slope.")
    bullet(doc, "The combined ratio, because it is the clearest evidence the operating turn is real. Down about four "
                "points through the first half of 2026 is the first data point, and the trend is what Best will "
                "weigh at the next review.")
    bullet(doc, "Whether Best still believes the plan. Its judgement that our growth turns to profit is what holds "
                "operating performance at Adequate. That judgement is doing real work in our rating, and it is "
                "renewed annually rather than granted.")
    bullet(doc, "The outlook, which is our earliest warning. A move from Stable to Negative signals that Best's view "
                "is shifting and gives us time to act before a letter changes.")
    body(doc, "On the story, the drawdown is easier to defend when it is presented as a planned, funded use of "
              "capital with a defined trough and a return path, rather than discovered by Best as an unexplained "
              "decline. The same applies to the operating turn: we should be putting the improving combined ratio "
              "in front of Best against the forecast it already holds, so each review confirms a trajectory we "
              "described in advance. Best has told us it expects us to maintain the Strongest assessment. The work "
              "is making the evidence arrive on schedule.")

    H(doc, "Summary")
    body(doc, "The rating is built from four blocks, and the balance sheet sets the floor under the other three. We "
              "hold the top balance-sheet grade on a BCAR of 67.3%, third of the 44 carriers we hold reports for, "
              "and none of the qualifiers that cost peers the top tier currently applies to us. Our plan draws that "
              "capital down through about 2029.")
    body(doc, "Two blocks can realistically move us. Each on its own costs one letter and lands us at A-, still "
              "inside the appetite. Both together cost two and take us to B++, outside it, though that case is a "
              "trough rather than an end state: returning operating performance to Adequate brings us back to A-, "
              "and it requires only that we stop losing a notch rather than earn one. The base case is that we hold "
              "A, and the first half of 2026 supports it. The exposure is that both live blocks are stressed in the "
              "same window, so the discipline is to keep the capital path explained in advance, the operating "
              "evidence arriving on time, and the preneed and reinsurance relationships intact through the trough.")

    note = doc.add_paragraph(); rn = note.add_run(
        "Sources and method. Wellabe's rating, grades and financials come from the AM Best Credit Report for "
        f"Wellabe Group, AMB #070369, effective May 2026, and prior reports. Peer grades and BCAR scores are taken "
        f"from the Best credit reports for {len(bcar_peers())} carriers. Statutory financials and RBC come from S&P "
        "Capital IQ and SNL; RBC is on the CAL basis. The RBC forward path and the likelihood language are "
        "illustrative, drawn from the 2026 strategic plan and our own judgement for discussion, not AM Best output. "
        "Internal and confidential, for ELT use.")
    rn.italic = True; rn.font.size = Pt(8.5); rn.font.color.rgb = RGBColor.from_string("8A8A8A")
    note.paragraph_format.space_before = Pt(14)
    doc.save(OUT)


if __name__ == "__main__":
    fig_bcar_history(); fig_earn_tiers(); fig_ladder()
    fig_bcar_tiers(); fig_bcar_vs_rbc()
    build()
    print("wrote " + str(OUT.relative_to(ROOT)))
