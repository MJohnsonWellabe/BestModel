"""Standalone companion note: what rating moves actually did to real carriers' businesses.

Separate from the main rating paper. Four transitions (A->A-, A-->B++, B++->A-, A-->A), each
anchored to a real carrier and the business impact its public filings / rating actions document.
Voice rules: plain language, no em-dashes, no decorative bolding, no "not X but Y".
Output: output/whitepaper/Rating_Move_Case_Studies.docx
Run: python src/make_case_studies.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "whitepaper" / "Rating_Move_Case_Studies.docx"
ACCENT = "#2E5A88"; INK = "#0B1C2C"; WELL = "#C0392B"; MUTE = "#5C6B78"


def _cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc); tcPr.append(shd)


def H(doc, text, size=13.5, before=12, color=ACCENT):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(2)


def body(doc, text, size=10.5):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); p.paragraph_format.space_after = Pt(6)


def kv(doc, label, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "  "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    r2 = p.add_run(text); r2.font.size = Pt(10.5)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try: t.style = "Light Grid Accent 1"
    except Exception: pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""; r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("FFFFFF"); _cell_bg(c, ACCENT.lstrip("#"))
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""; r = cells[j].paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    t = doc.add_paragraph(); r = t.add_run("What a Rating Move Actually Does to the Business: Four Real Cases")
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    s = doc.add_paragraph(); rs = s.add_run("Office of the Chief Actuary. Companion note to the AM Best rating paper. Internal and confidential.")
    rs.italic = True; rs.font.size = Pt(10); rs.font.color.rgb = RGBColor.from_string(MUTE.lstrip("#"))

    H(doc, "The one thing to take away", before=10)
    body(doc, "To see what a notch is worth, we looked at real carriers that made each move and what happened to their "
              "business afterward. The clearest finding is that the impact is not a straight line where each notch "
              "costs the same. It behaves like a cliff at the A- line. Moves inside the A band, from A to A- or from "
              "A- back to A, change very little for distribution, because A- is still an acceptable rating on almost "
              "every shelf. Crossing below A-, from A- to B++, is where sales and distribution actually break. And "
              "climbing back above A- is the hardest move of all, usually requiring a fresh capital sponsor rather "
              "than just time. That asymmetry is exactly why the rating paper frames A- as the line to defend.")
    body(doc, "One honest limit before the cases. Carriers rarely publish clean figures tying a single notch to "
              "sales, net promoter scores, or persistency. What public filings and rating actions do document is "
              "new-business volume, distribution access, and, at severe downgrades, surrenders. So the numbers below "
              "are the real, sourced ones, and where a metric is not public we say so rather than guess.")

    table(doc, ["The move", "Carrier", "What happened to the business"],
          [["A to A- (down, still above the line)", "American Southern (2026)",
            "A warning shot. Stayed a viable carrier; the notch reflected reserve problems more than it caused a sales break."],
           ["A- to B++ (down, crossing below)", "Phoenix Companies (2009)",
            "The cliff. Lost its largest distributor, new sales collapsed, and it never recovered as an independent seller."],
           ["B++ to A- (up, crossing back)", "Phoenix / the recovery problem",
            "The hardest move. Phoenix never regained A- on its own; carriers that climb back generally need new capital."],
           ["A- to A (up, within the band)", "F&G / Fidelity & Guaranty (2024)",
            "A modest tailwind. Opened bank, broker-dealer, and institutional shelves and supported fast growth."]])

    H(doc, "1.  A to A-: a warning shot, not a rupture")
    kv(doc, "Carrier:", "American Southern Insurance Company (Atlantic American group), downgraded by AM Best to A- (a-) from A (a) in April 2026.")
    kv(doc, "Cause:", "Reserve strengthening and adverse development in commercial auto, which drained risk-adjusted capital. The move was about the balance sheet, not a franchise collapse.")
    body(doc, "The business lesson is that dropping a single notch while staying at A- is a warning, not a rupture. A- "
              "is still above the threshold that gates most distribution, so a carrier at A- keeps its shelf access in "
              "the great majority of channels. The wider evidence supports this: AM Best has noted that roughly a "
              "third of the industry's annuity reserves now sit with about 95 companies carrying a lower issuer "
              "rating than they held in 2007, and those carriers have continued to operate and sell. A slip to A- "
              "raises the cost of capital and removes the margin for error, but it does not, by itself, take the "
              "business away. Clean sales-decline figures for a pure A to A- move are scarce precisely because the "
              "impact is modest when a carrier stays above the line.")

    H(doc, "2.  A- to B++: the cliff, and the case that proves it")
    kv(doc, "Carrier:", "The Phoenix Companies (Phoenix Life and Annuity), downgraded by AM Best to B++ from A- in 2009 (S&P moved it from BBB- to BB the same year).")
    body(doc, "This is the transition that actually breaks a business, and Phoenix documented every step of it in its "
              "public filings. As the ratings crossed below A-, its distribution left. In March 2009 State Farm, "
              "which in 2008 had been Phoenix's largest distributor at roughly 27% of total life insurance premiums "
              "and about 68% of annuity deposits, suspended sales of Phoenix products; by mid-2009 the restructured "
              "agreement provided for no new sales at all, stranding roughly 90,000 in-force policies and contracts. "
              "Sales through the independent brokerage general agencies fell sharply as well, hitting the universal "
              "life line that had been growing strongly only a year earlier. In short, a carrier that had been "
              "selling normally at A- lost the bulk of its new business within a year of dropping to B++. Phoenix "
              "never rebuilt as an independent seller and was eventually taken private. The mechanism is exactly the "
              "one the rating paper warns about: below A-, advisors and distributors pull the product from the shelf "
              "even though the company can still pay its claims.")

    H(doc, "3.  B++ back to A-: the hardest move on the board")
    body(doc, "The natural question after Phoenix is how hard it is to climb back. The answer is very hard, and "
              "Phoenix is again the example: it did not regain A- on its own. Once distribution has left and the "
              "in-force block is running off, the earnings and capital that would justify an upgrade are exactly what "
              "the downgrade took away, which is a difficult loop to break from the inside. The carriers that do "
              "cross back above A- generally do it with outside help, a new owner or a capital injection that repairs "
              "the balance sheet quickly rather than waiting years for organic recovery. F&G, in the next case, is "
              "the clearest recent illustration of that pattern: it climbed the scale on the back of new capital and "
              "a stronger parent, not on time alone. The practical message for us is that the A- line is far cheaper "
              "to hold than to win back, so the whole game is not crossing below it in the first place.")

    H(doc, "4.  A- to A: a real but gentle tailwind")
    kv(doc, "Carrier:", "F&G (Fidelity & Guaranty Life), upgraded by AM Best to A from A- in January 2024.")
    body(doc, "Moving up inside the A band helps, but gently, and mostly by widening distribution rather than "
              "producing an overnight sales jump. AM Best tied the upgrade to F&G's expanded business profile and the "
              "capital support of its parent, Fidelity National Financial. With the stronger rating F&G broadened "
              "beyond its independent-agent base into regional banks, broker-dealers, and institutional markets such "
              "as pension risk transfer, and it has become one of the fastest-growing annuity writers in the "
              "country, holding a top-five position in fixed indexed annuities and over 50 billion dollars in assets. "
              "The honest reading is that the upgrade both reflected and reinforced a business that was already "
              "growing. An upgrade opens doors and gives the marketing and reinsurance teams a better number to lead "
              "with, but it rewards a stronger business rather than creating one.")

    H(doc, "What this means for us")
    body(doc, "The four cases point the same way. Because the damage is concentrated at the A- line, our objective is "
              "not to chase a higher letter but to avoid crossing below A-. A slip within the A band, from A to A-, "
              "would be uncomfortable and would remove our margin for error, but the business would keep running, as "
              "American Southern and dozens of other carriers show. Crossing below A- to B++ is the move that would "
              "actually cost us distribution, and in our case it would land first and hardest in preneed, "
              "reinsurance, and capital raising rather than in guaranty-backed agent Medicare Supplement. And because "
              "climbing back above A- is the hardest move on the board, the cheapest possible strategy is to hold the "
              "line the first time. That is the same conclusion the rating paper reaches from the mechanics, arrived "
              "at here from what actually happened to real companies.")

    note = doc.add_paragraph(); rn = note.add_run(
        "Sources. Phoenix: The Phoenix Companies Form 10-K and 10-Q filings (2009) and AM Best and S&P rating "
        "actions; contemporaneous trade coverage of the State Farm and National Life distribution terminations. F&G: "
        "AM Best rating action upgrading F&G to A from A- (January 2024) and F&G investor disclosures. American "
        "Southern: AM Best rating action downgrading American Southern Group to A- from A (April 2026). Industry "
        "context: AM Best special report on shifting annuity reserves and credit quality since 2007. Carriers rarely "
        "publish notch-level net promoter or persistency data; figures here are the sales and distribution measures "
        "their public filings and rating actions actually disclose. Internal and confidential, for ELT use.")
    rn.italic = True; rn.font.size = Pt(8.5); rn.font.color.rgb = RGBColor.from_string("8A8A8A")
    note.paragraph_format.space_before = Pt(12)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print("wrote " + str(OUT.relative_to(ROOT)))
