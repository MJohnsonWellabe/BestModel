"""Build the ELT capital-appetite paper (.docx) - a fresh, human-voiced rewrite.

Audience: Executive Leadership Team, strategy-retreat capital-appetite session.
Self-contained: builds its own box-and-whisker figures and pulls a live peer table
from tool/data.json. Voice rules (per ELT feedback): plain language, no em-dashes,
no decorative bolding, no "not X but Y" constructions, no jargon-for-its-own-sake.
Output: output/whitepaper/Wellabe_AMBest_Rating_ELT.docx
Run: python src/make_elt_paper.py
"""
from __future__ import annotations
import json
import statistics as st
from pathlib import Path
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "output" / "whitepaper" / "figures"
FIG.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "output" / "whitepaper" / "Wellabe_AMBest_Rating_ELT.docx"
D = json.load(open(ROOT / "tool" / "data.json"))["carriers"]

ACCENT = "#2E5A88"; INK = "#0B1C2C"; WELL = "#C0392B"; MUTE = "#5C6B78"
TIERCOL = {"Strongest": "#0F5C8C", "Very Strong": "#27A35A", "Strong": "#9AA0A6", "Adequate": "#C97B2B"}
plt.rcParams.update({"font.size": 11, "font.family": "DejaVu Sans", "axes.edgecolor": "#C9BFA8",
                     "axes.titlesize": 13, "axes.titleweight": "bold", "figure.dpi": 150})


def W():
    return next(c for c in D if c.get("is_wellabe"))


# ----------------------------------------------------------------- figures
def fig_bcar():
    yrs = ["2022", "2023", "2024", "2025"]; bcar = [73.4, 73.0, 71.2, 67.3]
    fig, ax = plt.subplots(figsize=(7.0, 3.5), constrained_layout=True)
    ax.bar(yrs, bcar, color=ACCENT, width=0.62, zorder=3)
    for x, v in zip(yrs, bcar):
        ax.text(x, v + 1.6, f"{v:.0f}%", ha="center", fontweight="bold", color=INK)
    ax.axhspan(0, 25, color=WELL, alpha=0.08, zorder=0)
    ax.axhline(25, color=WELL, lw=1.4, ls="--", zorder=2)
    ax.text(3.4, 28, "25% is the line for the top tier", color=WELL, ha="right", fontsize=10)
    ax.set_ylim(0, 88); ax.set_ylabel("Capital to spare after a 1-in-250-year loss (%)")
    ax.set_title("Best's own capital measure (BCAR): a wide cushion, holding up")
    ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "elt_bcar.png"); plt.close(fig)


def _box_by_tier(ax, groups, vals_key, log=False):
    """Horizontal box-and-whisker, one box per group, returns y positions."""
    data = []
    for g in groups:
        ys = [c[vals_key] for c in D if c.get("_g") == g and c.get(vals_key) is not None]
        data.append(ys)
    pos = list(range(len(groups)))
    bp = ax.boxplot(data, positions=pos, vert=False, widths=0.55, patch_artist=True,
                    showfliers=False, medianprops=dict(color=INK, lw=1.6),
                    whiskerprops=dict(color="#8A8A8A"), capprops=dict(color="#8A8A8A"))
    for patch, g in zip(bp["boxes"], groups):
        patch.set_facecolor(TIERCOL.get(g, "#9AA0A6")); patch.set_alpha(0.35); patch.set_edgecolor("#8A8A8A")
    # scatter the underlying points lightly
    for i, ys in enumerate(data):
        ax.scatter(ys, [i] * len(ys), color=TIERCOL.get(groups[i], "#9AA0A6"), s=18, alpha=0.55,
                   zorder=3, edgecolor="white", lw=.4)
    ax.set_yticks(pos); ax.set_yticklabels(groups)
    if log:
        ax.set_xscale("log")
    return pos


def fig_rbc():
    order = ["Strongest", "Very Strong", "Strong", "Adequate"]
    for c in D:
        c["_g"] = c.get("bs_assessment")
    fig, ax = plt.subplots(figsize=(7.0, 3.7), constrained_layout=True)
    _box_by_tier(ax, order, "rbc_cal_pct", log=True)
    w = W()
    ax.scatter([w["rbc_cal_pct"]], [0], marker="D", s=170, color=WELL, edgecolor=INK, lw=1.4, zorder=6)
    ax.annotate("Wellabe, about 650%", (w["rbc_cal_pct"], 0), xytext=(6, 16), textcoords="offset points",
                color=WELL, fontweight="bold", fontsize=10)
    ax.set_xticks([200, 300, 500, 1000, 2000]); ax.set_xticklabels(["200", "300", "500", "1,000", "2,000"])
    ax.set_xlabel("NAIC RBC ratio, CAL basis (%)")
    ax.set_title("Capital by balance-sheet tier, and where we sit")
    ax.invert_yaxis(); ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "elt_rbc.png"); plt.close(fig)


def fig_op():
    order = ["Strong", "Adequate", "Marginal"]
    for c in D:
        c["_g"] = c.get("op_assessment")
    fig, ax = plt.subplots(figsize=(7.0, 3.7), constrained_layout=True)
    _box_by_tier(ax, order, "roe_5yr_mean")
    w = W()
    ax.scatter([w["roe_5yr_mean"]], [1], marker="D", s=170, color=WELL, edgecolor=INK, lw=1.4, zorder=6)
    ax.annotate("Wellabe: losing money, still held at Adequate", (w["roe_5yr_mean"], 1),
                xytext=(8, -30), textcoords="offset points", color=WELL, fontweight="bold", fontsize=9.5)
    for nm, lab in [("Pekin Life", "Pekin (A-)"), ("Government Personnel Mutual", "GPM (B++)")]:
        c = next((x for x in D if x["rating_unit_name"] == nm), None)
        if c and c.get("roe_5yr_mean") is not None:
            ax.annotate(lab, (c["roe_5yr_mean"], 2), xytext=(0, 10), textcoords="offset points",
                        ha="center", fontsize=9, color="#7A2418")
    ax.axvline(0, color="#C9BFA8", lw=1)
    ax.set_xlabel("Five-year average return on equity (%)")
    ax.set_title("Earnings by operating tier: we are at the bottom")
    ax.invert_yaxis(); ax.spines[["top", "right"]].set_visible(False)
    fig.savefig(FIG / "elt_op.png"); plt.close(fig)


def fig_ladder():
    fsr = ["A", "A-", "B++"]; xi = {f: i for i, f in enumerate(fsr)}
    rows = [
        ("Where we are today", "A", ACCENT, "Strongest capital, Adequate earnings"),
        ("Earnings slip one grade, to Marginal", "A-", "#27A35A", "our Strongest balance sheet holds the letter"),
        ("Capital falls two tiers, near 300% RBC", "A-", "#27A35A", "still A- if the franchise holds (see American Southern)"),
        ("Earnings fall further, or a second block slips", "B++", WELL, "takes two problems at once"),
    ]
    fig, ax = plt.subplots(figsize=(7.0, 3.6), constrained_layout=True)
    for i, (lab, f, col, note) in enumerate(rows):
        ax.barh(i, xi[f] + 0.5, color=col, alpha=.85, zorder=3, height=0.6)
        ax.text(xi[f] + 0.58, i, f, va="center", fontweight="bold", color=INK, fontsize=12)
        ax.text(xi[f] + 0.95, i, note, va="center", color=MUTE, fontsize=9)
        ax.text(-0.08, i, lab, va="center", ha="right", fontsize=10)
    ax.set_xlim(0, 5.2); ax.set_yticks([]); ax.invert_yaxis()
    ax.set_xticks(range(len(fsr))); ax.set_xticklabels(fsr)
    ax.set_title("How far we could fall, and what it takes")
    ax.set_xlabel("Financial strength rating")
    for s in ["top", "right", "left"]:
        ax.spines[s].set_visible(False)
    fig.savefig(FIG / "elt_ladder.png"); plt.close(fig)


# ----------------------------------------------------------------- docx helpers
def _cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def H(doc, text, size=14, before=12):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(ACCENT.lstrip("#"))
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, size=10.5):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6); return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(text); r.font.size = Pt(10.5); return p


def img(doc, name, caption):
    doc.add_picture(str(FIG / name), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTE.lstrip("#"))
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(10)


def table(doc, headers, rows, highlight=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("FFFFFF"); _cell_bg(c, ACCENT.lstrip("#"))
    for row in rows:
        cells = t.add_row().cells
        hot = highlight and row[0] == highlight
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
            if hot:
                r.bold = True; r.font.color.rgb = RGBColor.from_string(WELL.lstrip("#"))
                _cell_bg(cells[j], "FBE9E7")
    return t


# ----------------------------------------------------------------- peer table
CURATED = [
    "Guardian Life", "Physicians Mutual", "Mutual of Omaha", "Aflac",
    "Guarantee Trust Life", "National Guardian Life", "CNO — Bankers Life", "American National",
    "Globe Life", "American-Amicable / Trinity", "Wellabe Group",
    "National Western Life", "Assurity Life", "Pekin Life", "Funeral Directors Life",
    "Homesteaders Life", "Atlantic American — American Southern",
    "ManhattanLife — Assurance", "Government Personnel Mutual", "Investors Heritage",
    "ManhattanLife — Western United", "Continental General", "Sentinel Security Life",
]
FSR_ORDER = ["A++", "A+", "A", "A-", "B++", "B+", "B", "B-"]


def peer_rows():
    by = {c["rating_unit_name"]: c for c in D}
    out = []
    for nm in CURATED:
        c = by.get(nm)
        if not c:
            continue
        rbc = c.get("rbc_cal_pct")
        out.append([nm.replace("Wellabe Group", "Wellabe").replace(" — ", ", "),
                    c.get("fsr") or "n/a", c.get("bs_assessment") or "n/a",
                    c.get("op_assessment") or "n/a", c.get("bp_assessment") or "n/a",
                    f"{round(rbc):,}%" if rbc else "n/a"])
    out.sort(key=lambda r: (FSR_ORDER.index(r[1]) if r[1] in FSR_ORDER else 99,
                            -float(r[5].replace(",", "").rstrip("%")) if r[5] != "n/a" else 0))
    return out


# ----------------------------------------------------------------- build
def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

    t = doc.add_paragraph(); r = t.add_run("Our AM Best Rating: How It Works, Where We Stand, and What Would Move It")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    for line, sz in [("Office of the Chief Actuary. Strategy retreat, capital appetite session. ELT, internal and confidential.", 10),
                     ("Rated A (Excellent), Stable outlook. Wellabe Group, AMB #070369. Best's Credit Report effective May 2026.", 10)]:
        s = doc.add_paragraph(); rs = s.add_run(line); rs.italic = True; rs.font.size = Pt(sz)
        rs.font.color.rgb = RGBColor.from_string(MUTE.lstrip("#"))

    H(doc, "The short version", before=10)
    body(doc, "We are rated A (Excellent) with a Stable outlook. AM Best builds that letter from four things: our "
              "balance sheet, our earnings, our business franchise, and our risk management. We hold the highest "
              "possible mark on the balance sheet, and our capital sits in the top quartile of our peer "
              "group. We also hold the lowest letter inside the A band, and the reason is earnings. We have lost "
              "money three years running while we pay the up-front cost of fast Medicare Supplement growth, so our "
              "returns sit at the bottom of the group.")
    body(doc, "Here is the most useful thing to know for a conversation about capital. Our balance sheet is so far "
              "ahead of what the rating needs that capital is not where our rating risk comes from. We could give "
              "back a large part of our capital cushion and still hold a strong rating. The one thing that can "
              "actually move us is whether AM Best keeps believing our growth will turn into profit. That is an "
              "earnings and credibility question, and capital cannot buy our way out of it.")
    body(doc, "So the capital appetite question is not how much capital we need to protect the rating. It is how "
              "much of our cushion we are free to put to work, because the rating does not require us to carry as "
              "much as we do.")

    # 1. how it works
    H(doc, "1.  How AM Best builds the rating")
    body(doc, "The rating works off published tables, so we can read our own rating and any competitor's the same "
              "way. The balance sheet sets a starting point. The other three pieces move us up or down from there.")
    body(doc, "Step one. The balance-sheet grade sets a starting letter:")
    table(doc, ["Balance-sheet grade", "Starting point (ICR)", "Letter it implies"],
          [["Strongest", "a+ / a", "A"], ["Very Strong", "a / a-", "A / A-"],
           ["Strong", "a- / bbb+", "A- / B++"], ["Adequate", "bbb+ / bbb / bbb-", "B++ / B+"]],
          highlight="Strongest")
    body(doc, "Step two. Each of the other three pieces adds or removes notches:")
    table(doc, ["Building block", "How far it can move us", "The neutral, no-change grade"],
          [["Operating performance (earnings)", "up 2, down 3", "Adequate"],
           ["Business profile (franchise)", "up 2, down 2", "Neutral"],
           ["Risk management", "up 1, down 4", "Appropriate"]])
    body(doc, "Two things are worth noticing. The downside on each piece is bigger than the upside, so it is easier "
              "to lose ground than to gain it. And a great balance sheet only gets us to the starting line. A "
              "carrier can open at A on capital and still finish at B++ if its earnings and franchise are weak. The "
              "capital sets the floor under us. The other three blocks decide where we actually land.")

    # 2. capital
    H(doc, "2.  Our capital, and why the regulator's number can mislead")
    body(doc, "Capital is what drives the balance-sheet grade, and the balance-sheet grade is where we are "
              "strongest. The trap is using the wrong capital number. Two different measures get mixed up in these "
              "conversations.")
    body(doc, "The first is the NAIC risk-based capital ratio, the one we report on the CAL basis. We run about "
              "650%. That number is a regulatory floor and an early-warning gauge. Once a company is well above the "
              "floor, it does not rank healthy carriers against each other, and it is not the number AM Best uses "
              "to set a rating. Globe Life makes the point: it runs an RBC ratio around 316%, less than half of "
              "ours, and it is rated A, the same letter we hold. A 316% RBC ratio does not keep a company out of "
              "the A band. So we should be careful about reading too much into our own RBC number rising or falling.")
    body(doc, "The second measure is the one that matters for the rating, and it is Best's own capital model, "
              "called BCAR. Best takes our actual balance sheet and runs it through a series of bad years, up to "
              "roughly a 1-in-250-year loss. A company that still has capital left over after that loss earns the "
              "top grade. The bar for the top grade is having about 25% of capital still to spare. We have about "
              "67% to spare. We could give back a large part of that cushion and still clear the bar.")
    img(doc, "elt_bcar.png", "Figure 1. Best's BCAR cushion. The bar to be graded Strongest is 25%. We sit near 67% and it has held up while we fund growth.")
    body(doc, "One honest point about BCAR. It is a snapshot of today's balance sheet under stress, not a forecast. "
              "What looks ahead is Best's rating opinion and outlook, which reflect the losses Best already expects "
              "us to run. Our outlook is Stable while those expected losses are in front of us, which tells us Best "
              "is comfortable that the balance sheet stays strong through the plan.")

    # 3. the big finding
    H(doc, "3.  The most important thing about our capital")
    body(doc, "Because the balance-sheet grade comes from BCAR and not from the RBC ratio, our capital is much "
              "safer than the RBC number makes it look. Work through what a real decline would do.")
    body(doc, "Say our capital fell hard, down to something like 300% RBC. That is a big drop from 650%, and it "
              "would likely cost us two full balance-sheet tiers, from Strongest down to Strong. Even then, a Strong "
              "balance sheet starts us at a- or bbb+, which is A- or B++. As long as our franchise grade holds at "
              "Neutral and our earnings do not also collapse, that starting point lands us at A-. In plain terms, we "
              "could give back most of our capital cushion and most likely still hold A-. The only way capital "
              "pulls us below A- is if a second block breaks at the same time.")
    body(doc, "We do not have to guess at this. American Southern, a carrier in the Atlantic American group, runs a "
              "Strong balance sheet at about 213% RBC, well under half our capital, with the same Neutral franchise "
              "grade we hold, and AM Best rates it A-. It is a real company two tiers down from us on capital, below "
              "the 300% level, still holding A-.")
    body(doc, "That is the finding to carry into the capital appetite discussion. Our rating is not standing on a "
              "knife's edge of capital. There is a lot of room underneath us.")

    # 4. where we stand
    H(doc, "4.  Where we stand today")
    body(doc, "Best grades us Strongest on the balance sheet, Adequate on earnings, Neutral on franchise, and "
              "Appropriate on risk management. The balance sheet opens us at A, the other three net to no change, "
              "and we land at A.")
    body(doc, "The thing that shapes everything else is our recent earnings. We have run losses for three straight "
              "years, and the losses have grown each year. Our accident-and-health combined ratio has climbed every "
              "year, and surplus has come down since 2021.")
    table(doc, ["Year", "Net income", "Capital and surplus", "A&H combined ratio", "BCAR"],
          [["2021", "+$22M", "$630M", "95%", "n/a"], ["2022", "-$1M", "$615M", "99%", "73.4%"],
           ["2023", "-$21M", "$602M", "103%", "73.0%"], ["2024", "-$52M", "$560M", "109%", "71.2%"],
           ["2025", "-$71M", "$531M", "116%", "67.3%"]])
    body(doc, "We are still A and still Stable while running the largest losses in our history. The reason is that "
              "AM Best does not read these as ordinary losses. Statutory accounting makes us book the full cost of "
              "writing a new policy right away, so a record year of Medicare Supplement sales, with premium up 13% "
              "in 2025, shows up as a loss today on business we expect to earn back over its life. Best is giving us "
              "credit for that. Its own rating drivers say it plainly: it will act against us if losses run past "
              "what we projected, and it will only raise us if we show real, profitable growth, which it currently "
              "thinks is unlikely. So the rating rests on us delivering close to the plan we showed them.")
    body(doc, "The single number to watch is the combined ratio. It reached 116% in 2025, and the whole story turns "
              "on it bending back toward 100%. Early 2026 results are the first real test of whether that bend is "
              "starting. If the combined ratio turns this year, that is the clearest evidence we can give Best that "
              "the plan is working, and it is worth putting in front of them as soon as we are confident in it.")

    # 5. the peer landscape
    H(doc, "5.  The peer landscape")
    body(doc, "Reading the same four-block grades across the competitive set shows where the letters really come "
              "from. The table is sorted by rating, then by capital. The capital column and the rating do not move "
              "together.")
    table(doc, ["Carrier", "Rating", "Balance sheet", "Earnings", "Franchise", "RBC (CAL)"],
          peer_rows(), highlight="Wellabe")
    body(doc, "Three things stand out. Capital does not sort the ratings. Globe Life holds our same A on a weaker "
              "balance sheet and a 316% RBC ratio, carried by its size and franchise. Guarantee Trust Life holds an "
              "A above 800% RBC. ManhattanLife sits at B++ with more reported capital than several A- carriers. What "
              "sorts the ratings is earnings and franchise. And the senior-market specialists that look most like us "
              "cluster around A- and B++. We hold A out of that neighborhood on the strength of our balance sheet "
              "and the credit Best gives our growth plan.")
    img(doc, "elt_rbc.png", "Figure 2. Capital by balance-sheet tier across the peer set. We sit high. The point is how much capital each tier actually needs, not how much we happen to hold.")

    # 6. over / under index
    H(doc, "6.  Where we over-index and where we under-index")
    body(doc, "Against that peer set, we lean hard in two opposite directions, and both are on purpose.")
    bullet(doc, "We over-index on capital. Only a handful of carriers in the whole sample earn the Strongest "
                "balance-sheet grade, and we are one of them. Our 650% RBC ratio sits around the 75th percentile. "
                "Most carriers that share our A letter carry a weaker balance sheet than we do.")
    bullet(doc, "We under-index on earnings. Our five-year return on equity sits in the bottom handful of the "
                "entire sample, and Best grades us Adequate where most A-rated carriers are graded Strong. On "
                "earnings alone we look like the A- and B++ group, not the A group.")
    body(doc, "The typical A-rated carrier is the mirror image of us: a Very Strong balance sheet paired with Strong "
              "earnings. We have held our A with capital and a believable plan where others hold it with profits. "
              "Best leans on the weaker of the blocks, so our weaker block, earnings, is the one that governs our "
              "rating, and our stronger block, capital, is the one with room to give.")
    img(doc, "elt_op.png", "Figure 3. Earnings by operating tier. We run losses yet are held at Adequate. The grade is a judgement about our plan, not a reading of this year's number.")

    # 7. what moves us
    H(doc, "7.  What would actually move us")
    body(doc, "Capital is far above the line, risk management is steady, and the franchise grade moves slowly. That "
              "leaves earnings as the block that can realistically move. Best is holding us at Adequate on a "
              "forecast: losses that crest and then ease as the Medicare Supplement block matures. As long as we "
              "track that forecast, we stay at A. If losses run well past it, the earnings grade slips toward "
              "Marginal and the rating goes to A-.")
    body(doc, "Here our Strongest balance sheet does more than look good on paper. It buys us a notch of protection "
              "on earnings. If our earnings slip one grade to Marginal, the math takes us from A to A-, a single "
              "letter, because we start from the top tier. Carriers that start one tier lower, from a Very Strong "
              "balance sheet, land at B++ on that same slip. Government Personnel Mutual is exactly that case: Very "
              "Strong balance sheet, Marginal earnings, rated B++. Pekin has the very same grades as GPM and holds "
              "A-, but only because the committee gave it a one-notch lift. We would reach A- on the math itself. To "
              "fall to B++ ourselves, earnings would have to weaken well past Marginal, or a second block would have "
              "to give way at the same time.")
    img(doc, "elt_ladder.png", "Figure 4. How far we could fall and what it takes. One block slipping costs one letter. B++ needs two things to go wrong at once.")
    body(doc, "Two features make this manageable. We usually get warning. A downgrade from A is almost always "
              "preceded by a move from a Stable to a Negative outlook, which tends to give a year or more of lead "
              "time. Our outlook is Stable today. And the move is reversible, because the earnings grade follows a "
              "trend, so a return to profit can win the letter back. American Southern, a sister company inside "
              "Atlantic American, is a recent reminder of how downgrades actually arrive: it came from reserve "
              "problems that drained capital, a specific event, not capital quietly drifting lower.")

    # 8. what we are watching elsewhere
    H(doc, "8.  The two slower-moving blocks")
    body(doc, "Franchise. Best grades us Neutral, and that is our ceiling for now. We are genuinely diversified, "
              "across Medicare Supplement, preneed, ancillary health, and life, across 40-plus states and several "
              "distribution channels, and that diversification is good and is what holds us at Neutral. The honest "
              "limit is that almost all of it still sits inside the senior market. To earn a better franchise grade "
              "we would need to grow well beyond that niche. Adding another senior-market product helps earnings and "
              "helps hold Neutral, but it does not lift the grade. One thing we cannot lose sight of: most of our "
              "Medicare Supplement sales come from our top 20 states, so a regulatory or competitive hit to Medicare "
              "Supplement would land on our earnings and our franchise at the same time. That is the one event that "
              "could move two blocks together.")
    body(doc, "Risk management. Best grades us Appropriate, and the goal is simply to keep it there. The real risk "
              "on this side is not the grade. It is letting a preventable problem, a bad reinsurance treaty, a "
              "reserve that proves short, an investment that goes wrong, grow into something that shows up in our "
              "capital or earnings. The grade is just the label on how well we do that work.")

    # 9. how it likely plays out
    H(doc, "9.  How this likely plays out")
    body(doc, "These are judgements for discussion, not model output. They reflect the plan as we understand it, "
              "where surplus troughs around 2028 and Medicare Supplement earnings turn positive in 2029.")
    table(doc, ["Path", "Rough odds", "What it looks like"],
          [["Plan holds", "about 55%", "Combined ratio turns, we stay A and Stable through 2030."],
           ["A scare, no downgrade", "about 25%", "Outlook goes Negative around 2027 or 2028, then back to Stable by 2029 or 2030 as earnings turn."],
           ["A one-letter dip", "about 15%", "If the turn comes late, a move to A- around 2028 or 2029, working back toward A by the early 2030s."],
           ["Below A-", "about 5%", "Only if a second block breaks alongside earnings. Needs two things to go wrong."]])
    body(doc, "The shape of this is the reassuring part. The most likely paths keep us at A or, at worst, a "
              "recoverable A-, and we should see a Negative outlook coming before any downgrade. We would treat a "
              "move to a Negative outlook as the signal to act, and we would bring it to the board right away.")

    # 10. capital appetite
    H(doc, "10.  What this means for capital appetite")
    body(doc, "Pulling it together for the capital question in front of us:")
    bullet(doc, "Set appetite by the plan, not by the rating. The capital the rating needs is well below what we "
                "carry. The binding use of capital is funding new-business strain through the 2028 trough with a "
                "margin for a bad year, not defending the letter A.")
    bullet(doc, "Our cushion is available to put to work, and using it well is what protects the rating, because the "
                "rating is protected by the earnings our growth produces, not by surplus sitting idle. Holding extra "
                "capital we do not need buys us no rating benefit.")
    bullet(doc, "Do not over-react to a falling RBC number. Even a drop to around 300% most likely keeps us at A-, "
                "and BCAR, the measure that counts, has far more room than the RBC ratio suggests.")
    bullet(doc, "Keep the unused flexibility in view. We have no debt, no surplus notes outstanding yet, FHLB Des "
                "Moines capacity, and live reinsurance options. Best already counts these toward the balance sheet, "
                "so they are real appetite headroom that the RBC ratio does not show.")
    body(doc, "On the surplus note we are working toward: it raises capital and helps both the RBC ratio and BCAR, "
              "and Best treats access to it as a sign of financial flexibility, which is a positive. The trade-off "
              "is that a surplus note is slightly lower-quality capital than retained earnings and carries an "
              "interest cost, so it gives a small offset on the capital-quality side. On balance it helps the "
              "cushion and does not change the rating story. It is funding for the growth plan, not a fix for it.")
    body(doc, "One caveat for honesty. Balance-sheet strength is more than the headline ratio. On Best's reported "
              "measure, our higher-risk asset leverage rose to roughly 40% of capital and surplus in 2025, up from "
              "about 29% in 2023. That is no threat to the Strongest grade today, but it is the kind of drift to "
              "keep an eye on as we lean on the cushion.")

    # 11. takeaways
    H(doc, "What to take away")
    bullet(doc, "We are rated A, Stable, and we expect to hold A- or better through at least 2030.")
    bullet(doc, "Capital is our strongest block and is not where our rating risk comes from. We carry far more than the rating needs, and even a large drop most likely leaves us at A-.")
    bullet(doc, "Earnings are the one block that can really move us. We are at the bottom of the peer group on returns, and the rating rests on Best believing our growth turns to profit.")
    bullet(doc, "The combined ratio is the number to watch. If it turns in 2026, that is the best evidence we can show Best that the plan is working.")
    bullet(doc, "Capital appetite can be set by the growth plan, not by the rating, because the cushion is there to fund the bet that protects the letter.")

    note = doc.add_paragraph(); rn = note.add_run(
        "Sources and method. Wellabe's rating, grades, and financials come from the AM Best Credit Report for "
        "Wellabe Group, AMB #070369, effective May 2026, and prior reports. Peer figures come from a 50-carrier "
        "model built from public AM Best grades and S&P Capital IQ and SNL statutory data. RBC is on the CAL basis "
        "(total adjusted capital divided by company action level). Peer figures are sample approximations, not AM "
        "Best's internal composites. The likelihood ranges in section 9 are our own judgement for discussion. "
        "Internal and confidential, for ELT use.")
    rn.italic = True; rn.font.size = Pt(8.5); rn.font.color.rgb = RGBColor.from_string("8A8A8A")
    note.paragraph_format.space_before = Pt(14)
    doc.save(OUT)


if __name__ == "__main__":
    fig_bcar(); fig_rbc(); fig_op(); fig_ladder()
    build()
    print("wrote " + str(OUT.relative_to(ROOT)))
