"""Generate board-ready figures (PNG) and assemble the Wellabe rating white paper as a .docx.

Figures come from the 50-carrier peer model (tool/data.json) + the BCAR series read from the
four Best's Credit Reports. Output: output/whitepaper/figures/*.png and
output/whitepaper/Wellabe_AMBest_Rating_Board.docx.
Run: python src/make_board_doc.py
"""
from __future__ import annotations
import json
from pathlib import Path
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "output" / "whitepaper" / "figures"
FIG.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "output" / "whitepaper" / "Wellabe_AMBest_Rating_Board.docx"
D = json.load(open(ROOT / "tool" / "data.json"))["carriers"]

INK = "#0B1C2C"; ACCENT = "#2E5A88"; WELL = "#C0392B"
RAMP = {"Strongest": "#0F5C8C", "Very Strong": "#27A35A", "Strong": "#9AA0A6", "Adequate": "#F2A93B"}
plt.rcParams.update({"font.size": 12, "font.family": "DejaVu Sans", "axes.edgecolor": "#C9BFA8",
                     "axes.titlesize": 14, "axes.titleweight": "bold", "figure.dpi": 150})


def fig_bcar():
    yrs = ["YE2022", "YE2023", "YE2024", "YE2025"]; bcar = [73.4, 73.0, 71.2, 67.3]
    fig, ax = plt.subplots(figsize=(7.2, 3.4))
    ax.bar(yrs, bcar, color=ACCENT, width=0.6, zorder=3)
    for x, v in zip(yrs, bcar):
        ax.text(x, v + 1.5, f"{v:.1f}%", ha="center", fontweight="bold", color=INK)
    ax.axhspan(0, 25, color="#C0392B", alpha=0.10, zorder=0)
    ax.axhline(25, color=WELL, lw=1.5, ls="--", zorder=2)
    ax.text(3.45, 27, "25% = bar to be assessed “Strongest”", color=WELL, ha="right", fontsize=11)
    ax.set_ylim(0, 90); ax.set_ylabel("BCAR at 99.6% VaR (%)")
    ax.set_title("Risk-adjusted capital (BCAR): a wide, stable cushion")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout(); fig.savefig(FIG / "fig1_bcar.png"); plt.close(fig)


def fig_rbc():
    order = ["Strongest", "Very Strong", "Strong", "Adequate"]
    data = {t: [c["rbc_cal_pct"] for c in D if c.get("bs_assessment") == t and c.get("rbc_cal_pct")] for t in order}
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    for i, t in enumerate(order):
        ys = data[t]
        ax.scatter(ys, [i] * len(ys), color=RAMP[t], s=42, alpha=0.85, zorder=3, edgecolor="white", lw=.5)
        if ys:
            import statistics as st
            ax.scatter([st.median(ys)], [i], marker="|", s=900, color=INK, zorder=4)
    well = next((c for c in D if c.get("is_wellabe") and c.get("rbc_cal_pct")), None)
    if well:
        ax.scatter([well["rbc_cal_pct"]], [0], marker="D", s=180, color=WELL, edgecolor=INK, lw=1.5, zorder=5)
        ax.annotate("Wellabe ~648%", (well["rbc_cal_pct"], 0), xytext=(8, 14), textcoords="offset points",
                    color=WELL, fontweight="bold")
    for x, lab in [(533, "Strongest\nfloor ~530%"), (376, "Very Strong\n~375%"), (277, "Strong\n~275%")]:
        ax.axvline(x, color="#C9BFA8", ls=":", lw=1)
    ax.set_yticks(range(len(order))); ax.set_yticklabels(order)
    ax.set_xscale("log"); ax.set_xlabel("RBC ratio — CAL basis (%, log)")
    ax.set_xticks([200, 300, 500, 1000, 2000]); ax.set_xticklabels(["200", "300", "500", "1,000", "2,000"])
    ax.set_title("Where Wellabe sits on capital — and the tier floors")
    ax.invert_yaxis(); ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout(); fig.savefig(FIG / "fig2_rbc.png"); plt.close(fig)


def fig_op():
    order = ["Strong", "Adequate", "Marginal"]
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    for i, t in enumerate(order):
        ys = [c["roe_5yr_mean"] for c in D if c.get("op_assessment") == t and c.get("roe_5yr_mean") is not None]
        col = {"Strong": "#27A35A", "Adequate": "#F2A93B", "Marginal": "#C0392B"}[t]
        ax.scatter(ys, [i] * len(ys), color=col, s=42, alpha=0.8, zorder=3, edgecolor="white", lw=.5)
    label = {"Lumico (RGA)": "Lumico", "Government Personnel Mutual": "GPM",
             "American-Amicable / Trinity": "Am-Amicable"}
    for c in D:
        if c["rating_unit_name"] in label and c.get("roe_5yr_mean") is not None:
            ax.annotate(label[c["rating_unit_name"]], (c["roe_5yr_mean"], 2), xytext=(0, 9),
                        textcoords="offset points", ha="center", fontsize=9, color="#7A2418")
    well = next((c for c in D if c.get("is_wellabe") and c.get("roe_5yr_mean") is not None), None)
    if well:
        ax.scatter([well["roe_5yr_mean"]], [1], marker="D", s=180, color=WELL, edgecolor=INK, lw=1.5, zorder=5)
        ax.annotate("Wellabe — held Adequate\ndespite negative ROE", (well["roe_5yr_mean"], 1),
                    xytext=(10, -28), textcoords="offset points", color=WELL, fontweight="bold", fontsize=10)
    ax.axvline(0, color="#C9BFA8", lw=1)
    ax.set_yticks(range(len(order))); ax.set_yticklabels(["Strong", "Adequate", "Marginal"])
    ax.set_xlabel("5-year mean ROE (%)"); ax.invert_yaxis()
    ax.set_title("Operating performance: a belief call, not just a number")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout(); fig.savefig(FIG / "fig3_operating.png"); plt.close(fig)


def fig_scenarios():
    fsr = ["A++", "A+", "A", "A-", "B++", "B+", "B"]; yi = {f: i for i, f in enumerate(fsr)}
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    rows = [("Today\nStrongest / Adequate", "A", ACCENT, ""),
            ("Balance sheet −1 tier\nVery Strong / Adequate", "A-", "#27A35A", "(e.g. Pekin, FDLIC)"),
            ("Operating −1\nStrongest / Marginal", "A-", "#27A35A", "(capital cushions it)"),
            ("Both slip\nVery Strong / Marginal", "B++", WELL, "(exactly GPM today)")]
    for i, (lab, f, col, note) in enumerate(rows):
        ax.barh(i, yi[f] + 0.6, color=col, alpha=.85, zorder=3, height=0.6)
        ax.text(yi[f] + 0.65, i, f"{f}  {note}", va="center", fontweight="bold", color=INK, fontsize=11)
        ax.text(-0.1, i, lab, va="center", ha="right", fontsize=10)
    ax.set_xlim(0, 7); ax.set_xticks(range(len(fsr))); ax.set_xticklabels(fsr)
    ax.set_yticks([]); ax.invert_yaxis()
    ax.set_title("The downside ladder (predicted FSR)")
    ax.set_xlabel("Financial Strength Rating")
    for s in ["top", "right", "left"]:
        ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(FIG / "fig4_scenarios.png"); plt.close(fig)


# ------------------------------------------------------------------ docx assembly
def H(doc, text, size=15, color=ACCENT, space_before=10):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color.lstrip("#")); p.paragraph_format.space_before = Pt(space_before)
    return p


def body(doc, text):
    p = doc.add_paragraph();
    for chunk, bold in _md_bold(text):
        r = p.add_run(chunk); r.bold = bold; r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6); return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    for chunk, bold in _md_bold(text):
        r = p.add_run(chunk); r.bold = bold; r.font.size = Pt(10.5)
    return p


def _md_bold(text):
    out = []; parts = text.split("**")
    for i, part in enumerate(parts):
        if part: out.append((part, i % 2 == 1))
    return out


def img(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("5C6B78"); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)


def build_doc():
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    t = doc.add_paragraph(); r = t.add_run("Our AM Best Rating: Where We Stand and What Moves It")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor.from_string("0B1C2C")
    sub = doc.add_paragraph(); rs = sub.add_run(
        "Office of the Chief Actuary  ·  Board briefing, June 2026  ·  Rated A (Excellent), Stable outlook (AMB #070369)")
    rs.italic = True; rs.font.size = Pt(10); rs.font.color.rgb = RGBColor.from_string("5C6B78")

    H(doc, "Bottom line", 15)
    body(doc, "This briefing reverse-engineers how AM Best builds our rating, shows where we sit against 50 peer carriers, and isolates the one variable that actually moves us. No board action is required; the purpose is shared visibility on the rating and how we are managing it.")
    bullet(doc, "We are rated **A (Excellent)** with a **Stable** outlook, and we expect to hold **A- or better through at least 2027–2028**, the peak-strain window.")
    bullet(doc, "**Capital is not the constraint.** AM Best rates capital on its own model (BCAR), where we sit at **67% against a 25% bar** for the top tier — a ~2.7× cushion that has barely moved through three years of growth strain. Best explicitly expects us to keep the Strongest assessment.")
    bullet(doc, "**The one risk that matters is operating-performance credibility** — whether Best continues to believe our growth plan. We are held a full tier above look-alike carriers on that belief alone, so it is the variable we manage most deliberately.")
    bullet(doc, "**Realistic downside is A-, not lower.** A one-block slip lands A-; only simultaneous capital *and* earnings deterioration reaches B++ — and any move is pre-signalled by a shift to Negative outlook, which has not occurred. We will flag that signal to the board the moment it appears.")

    H(doc, "1.  Capital: a wide, stable cushion")
    body(doc, "Best does not use the regulator's RBC ratio to set our rating; it runs its own capital model (BCAR). Our BCAR has moved only 73% to 67% over four annual reviews despite funding rapid growth — and the bar for the top 'Strongest' tier is just 25%. We could lose more than half our risk-adjusted capital and still clear it. Best's published outlook says it expects Strongest to be maintained.")
    img(doc, FIG / "fig1_bcar.png", "Figure 1. BCAR (Best's capital measure) sits ~2.7x the Strongest threshold and has been stable through the growth phase.")
    body(doc, "On the regulator's RBC ratio — a useful internal proxy — we hold ~648%, inside the Strongest band. The peer floors (below) match our internal appetite: ~530% to hold Strongest, ~375% for the A- zone, ~275% for Strong. Even our plan's 400% RBC trough would most likely still read Strongest on BCAR. **Capital flexibility is an unused strength: no debt, no surplus notes, FHLB capacity, and live reinsurance options — all sources Best already credits.**")
    img(doc, FIG / "fig2_rbc.png", "Figure 2. Wellabe (red diamond) versus the peer capital distribution by balance-sheet tier; dotted lines mark the tier floors.")

    H(doc, "2.  Operating performance: the binding risk")
    body(doc, "This is where a downgrade would start. Carriers with earnings similar to or worse than ours — Lumico, GPM, American-Amicable — are assessed Marginal. We are held a tier higher at Adequate because Best reads our losses as **credited new-business strain in support of a growth plan it believes**. That belief, not the numbers, is the difference.")
    img(doc, FIG / "fig3_operating.png", "Figure 3. We run negative recent ROE yet are held at Adequate; the Marginal carriers (labelled) earn the same or more. The assessment is a judgement on our growth story.")
    body(doc, "**The cautionary example is Government Personnel Mutual: a Very Strong balance sheet with plenty of capital, yet rated B++ because its operating performance is Marginal.** Capital does not rescue a Marginal earnings assessment. The lever for us is therefore as much about evidence and communication as results: Best must be able to see the turn coming — rerates achieved (13.4% in 2025), the improved 2024-charter economics, the projected Med Supp earnings inflection in 2029, and positive strain-adjusted earnings in every year after 2025.")

    H(doc, "3.  Business profile: it is breadth, not size")
    body(doc, "Our business profile is Neutral and unlikely to change. Importantly, what earns Neutral is **not scale** — we sit in a cluster with Funeral Directors, Assurity and Sentinel, all assessed Limited and all our size or larger. What separates us is **breadth**: multi-line (Med Supp, preneed, ancillary health, life), a 40-plus-state footprint, broad distribution, and a direct-to-consumer channel. The story to tell Best is diversification and pricing sophistication, not size. The watch item is concentration: most Med Supp sales come from our top 20 states, and a regulatory or competitive shock to Med Supp would pressure earnings and franchise together.")

    H(doc, "4.  The downside ladder")
    body(doc, "Running our blocks through Best's published rating math gives a clear, bounded set of outcomes:")
    img(doc, FIG / "fig4_scenarios.png", "Figure 4. Predicted rating under each scenario, with the real carrier that proves each rung.")
    body(doc, "A single slip — capital or earnings — lands A-. Only both together reach B++, and the capital leg is unlikely given the BCAR cushion. The credibility test peaks in 2027–2028 (when growth strain peaks and surplus troughs before the 2029 earnings turn). Because the operating trigger is a *trend*, a return to positive earnings would restore the notch — the path back from B++ to A- is one good print, which is part of why Best tolerates the strain today.")

    H(doc, "5.  What we are leaning into")
    bullet(doc, "**Capital:** lead with BCAR, not RBC; keep surplus-note / reinsurance / FHLB flexibility visible and disciplined.")
    bullet(doc, "**Operating:** make the plan-versus-actual turn easy for Best to verify — rerate achievement, 2024-charter IRR, underwritten mix, persistency, the 2029 inflection.")
    bullet(doc, "**Business profile:** protect and publicise breadth and pricing sophistication; keep de-concentrating beyond Med Supp.")

    H(doc, "6.  What could change our view")
    body(doc, "A move to Negative outlook would be the first signal — we have none. The genuine triggers are (i) earnings materially below plan such that Best stops crediting the growth story, or (ii) a Med Supp shock hitting earnings and franchise at once. The balance-sheet trigger that the raw RBC number might suggest is, in practice, remote.")

    note = doc.add_paragraph(); rn = note.add_run(
        "Method: figures derived from a 50-carrier peer model built from public AM Best assessments and "
        "S&P Capital IQ / SNL statutory data, plus the BCAR series from our 2023–2026 Best's Credit Reports "
        "and the 2026 strategic plan. Peer-relative metrics are sample approximations, not AM Best's internal "
        "composites. Confidential — internal board use.")
    rn.italic = True; rn.font.size = Pt(8.5); rn.font.color.rgb = RGBColor.from_string("8A8A8A")
    note.paragraph_format.space_before = Pt(14)
    doc.save(OUT)


if __name__ == "__main__":
    fig_bcar(); fig_rbc(); fig_op(); fig_scenarios()
    build_doc()
    print("wrote figures + " + str(OUT.relative_to(ROOT)))
